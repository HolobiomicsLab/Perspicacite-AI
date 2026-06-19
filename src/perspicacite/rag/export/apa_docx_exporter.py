"""Export an agentic answer + its cited papers to an APA-formatted ``.docx``.

Opt-in via ``config.rag_modes.agentic.export_apa_docx`` (default off). Requires
the optional ``[docx]`` extra (``python-docx``), imported lazily so the
dependency is only needed when the feature is actually enabled.

Inputs are tolerant: ``papers`` may be ``Paper`` model instances, plain dicts,
or a mix. Authors may be ``Author`` objects (``.name``), dicts, or strings.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any


def _author_name(author: Any) -> str:
    """Best-effort author display name from an Author object / dict / string."""
    if isinstance(author, str):
        return author.strip()
    name = getattr(author, "name", None)
    if name is None and isinstance(author, dict):
        name = author.get("name")
    return str(name).strip() if name else ""


def format_authors(authors: list[Any] | None) -> str:
    """APA author list: ``A``, ``A & B``, or ``A, B, & C``."""
    names = [n for n in (_author_name(a) for a in (authors or [])) if n]
    if not names:
        return ""
    if len(names) == 1:
        return names[0]
    if len(names) == 2:
        return f"{names[0]} & {names[1]}"
    return ", ".join(names[:-1]) + f", & {names[-1]}"


def _field(paper: Any, name: str, default: Any = None) -> Any:
    """Read ``name`` from a Paper object or a dict."""
    if isinstance(paper, dict):
        return paper.get(name, default)
    return getattr(paper, name, default)


def paper_to_apa(paper: Any) -> str:
    """Render a single APA-style reference string from a Paper / dict."""
    authors = format_authors(_field(paper, "authors", []) or [])
    year = _field(paper, "year") or "n.d."
    title = (_field(paper, "title", "") or "").strip()
    journal = (_field(paper, "journal", "") or "").strip()
    doi = (_field(paper, "doi", "") or "").strip()

    text = f"{authors} ({year}). {title}.".strip()
    if journal:
        text += f" {journal}."
    if doi:
        doi_clean = doi.replace("https://doi.org/", "").strip()
        if doi_clean:
            text += f" https://doi.org/{doi_clean}"
    return text.strip()


def export_apa_docx(
    manuscript_text: str,
    papers: list[Any] | None,
    output_path: str | Path,
) -> Path:
    """Write *manuscript_text* + an APA reference list to a ``.docx``.

    Returns the written path. Lazily imports ``python-docx`` and raises a clear
    ``ImportError`` (with the install hint) when the optional extra is absent.
    Deduplicates references by their rendered string, preserving order.
    """
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - only without the [docx] extra
        raise ImportError(
            "APA .docx export requires the optional [docx] extra: "
            "`uv sync --extra docx` (installs python-docx)."
        ) from exc

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading("Manuscript", level=1)
    doc.add_paragraph(manuscript_text or "")

    doc.add_heading("References", level=1)
    seen: set[str] = set()
    for paper in papers or []:
        ref = paper_to_apa(paper)
        if ref and ref not in seen:
            seen.add(ref)
            doc.add_paragraph(ref)

    doc.save(str(out))
    return out
